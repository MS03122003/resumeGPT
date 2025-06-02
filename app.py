from flask import Flask, request, jsonify, session, render_template
from flask_cors import CORS
import os
import tempfile
import logging
from datetime import datetime
import uuid
import asyncio
from concurrent.futures import ThreadPoolExecutor
import json
import re
from typing import List, Dict, Any, Optional
import mimetypes

# Google APIs
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import Flow
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload, MediaFileUpload
from google.auth.transport.requests import Request
import google.generativeai as genai

# Document processing
from docx import Document
from docx.shared import Inches
import PyPDF2
import io
from pdf2docx import Converter
import mammoth

# Configuration
app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY')
CORS(app)

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Configuration
UPLOAD_FOLDER = 'uploads'
PROCESSED_FOLDER = 'processed'
SCOPES = [
    'https://www.googleapis.com/auth/drive',
    'https://www.googleapis.com/auth/documents'
]

# Create directories
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(PROCESSED_FOLDER, exist_ok=True)

# Initialize Gemini
genai.configure(api_key=os.environ.get('GEMINI_API_KEY'))

class DocumentProcessor:
    """Handles document conversion and processing"""
    
    @staticmethod
    def pdf_to_docx(pdf_path: str, docx_path: str) -> bool:
        """Convert PDF to DOCX format"""
        try:
            cv = Converter(pdf_path)
            cv.convert(docx_path, start=0, end=None)
            cv.close()
            return True
        except Exception as e:
            logger.error(f"Error converting PDF to DOCX: {e}")
            return False
    
    @staticmethod
    def extract_text_from_pdf(pdf_path: str) -> str:
        """Extract text from PDF as fallback"""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            return ""
    
    @staticmethod
    def extract_text_from_docx(docx_path: str) -> str:
        """Extract text from DOCX"""
        try:
            doc = Document(docx_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            return ""
    
    @staticmethod
    def create_combined_document(texts: List[str], output_path: str) -> bool:
        """Combine multiple texts into a single DOCX document"""
        try:
            doc = Document()
            doc.add_heading('Combined Resume Analysis', 0)
            
            for i, text in enumerate(texts, 1):
                doc.add_heading(f'Resume {i}', level=1)
                doc.add_paragraph(text)
                doc.add_page_break()
            
            doc.save(output_path)
            return True
        except Exception as e:
            logger.error(f"Error creating combined document: {e}")
            return False

class GoogleDriveHandler:
    """Handles Google Drive operations"""
    
    def __init__(self):
        self.credentials = None
        self.drive_service = None
        self.docs_service = None
    
    def authenticate(self, credentials_json: str) -> bool:
        """Authenticate with Google Drive"""
        try:
            flow = Flow.from_client_config(
                json.loads(credentials_json),
                scopes=SCOPES
            )
            flow.redirect_uri = 'http://localhost:8080/callback'
            
            auth_url, _ = flow.authorization_url(prompt='consent')
            return auth_url
        except Exception as e:
            logger.error(f"Authentication error: {e}")
            return False
    
    def build_services(self, credentials):
        """Build Google API services"""
        self.credentials = credentials
        self.drive_service = build('drive', 'v3', credentials=credentials)
        self.docs_service = build('docs', 'v1', credentials=credentials)
    
    def list_files_from_folder(self, folder_id: str) -> List[Dict]:
        """List all files from a Google Drive folder"""
        try:
            query = f"'{folder_id}' in parents and (mimeType='application/pdf' or mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document')"
            results = self.drive_service.files().list(
                q=query,
                fields="files(id, name, mimeType, size)"
            ).execute()
            
            return results.get('files', [])
        except Exception as e:
            logger.error(f"Error listing files: {e}")
            return []
    
    def download_file(self, file_id: str, file_name: str) -> str:
        """Download file from Google Drive"""
        try:
            request = self.drive_service.files().get_media(fileId=file_id)
            file_path = os.path.join(UPLOAD_FOLDER, file_name)
            
            with open(file_path, 'wb') as fh:
                downloader = MediaIoBaseDownload(fh, request)
                done = False
                while done is False:
                    status, done = downloader.next_chunk()
            
            return file_path
        except Exception as e:
            logger.error(f"Error downloading file: {e}")
            return None
    
    def convert_to_docs(self, file_id: str) -> str:
        """Convert file to Google Docs format"""
        try:
            # Export as Google Docs
            request = self.drive_service.files().export_media(
                fileId=file_id,
                mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            
            file_name = f"converted_{uuid.uuid4().hex}.docx"
            file_path = os.path.join(PROCESSED_FOLDER, file_name)
            
            with open(file_path, 'wb') as fh:
                downloader = MediaIoBaseDownload(fh, request)
                done = False
                while done is False:
                    status, done = downloader.next_chunk()
            
            return file_path
        except Exception as e:
            logger.error(f"Error converting to docs: {e}")
            return None

class GeminiAnalyzer:
    """Handles AI analysis using Gemini API"""
    
    def __init__(self):
        self.model = genai.GenerativeModel("gemini-2.0-flash")
    
    def analyze_resume(self, text: str, query: str) -> str:
        """Analyze resume text using Gemini"""
        try:
            prompt = f"""
            You are IMEJOBCV, an intelligent resume screening system. Analyze the following resume(s) and answer the user's query accurately and professionally.

            Resume Content:
            {text}

            User Query: {query}

            Instructions:
            1. Provide detailed and accurate analysis based on the resume content
            2. If looking for specific skills/experience, highlight relevant sections
            3. If comparing candidates, provide comparative analysis
            4. Be specific about years of experience, technologies, and achievements
            5. If the query cannot be answered from the resume, clearly state what information is missing

            Response:
            """
            
            response = self.model.generate_content(prompt)
            return response.text
        except Exception as e:
            logger.error(f"Error in Gemini analysis: {e}")
            return "I apologize, but I encountered an error while analyzing the resume. Please try again."
    
    def extract_resume_summary(self, text: str) -> Dict:
        """Extract structured summary from resume"""
        try:
            prompt = f"""
            Extract the following information from this resume and return as JSON:
            
            Resume: {text}
            
            Extract:
            - Name
            - Email
            - Phone
            - Years of Experience (estimate)
            - Key Skills (list)
            - Education
            - Current/Latest Job Title
            - Companies Worked At (list)
            - Key Achievements (list)
            
            Return only valid JSON format.
            """
            
            response = self.model.generate_content(prompt)
            try:
                return json.loads(response.text)
            except:
                return {"summary": response.text}
        except Exception as e:
            logger.error(f"Error extracting summary: {e}")
            return {"error": "Could not extract summary"}

# Initialize components
doc_processor = DocumentProcessor()
drive_handler = GoogleDriveHandler()
gemini_analyzer = GeminiAnalyzer()

# Store for processed documents
document_store = {}
@app.route('/')
def index():
    return render_template("index.html")
@app.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({"status": "healthy", "timestamp": datetime.now().isoformat()})

@app.route('/upload', methods=['POST'])
def upload_file():
    """Handle file upload"""
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file provided"}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({"error": "No file selected"}), 400
        
        # Save uploaded file
        filename = f"{uuid.uuid4().hex}_{file.filename}"
        file_path = os.path.join(UPLOAD_FOLDER, filename)
        file.save(file_path)
        
        # Process file
        processed_content = process_single_file(file_path)
        
        if processed_content:
            doc_id = str(uuid.uuid4())
            document_store[doc_id] = {
                "content": processed_content,
                "filename": file.filename,
                "upload_time": datetime.now().isoformat(),
                "type": "single_file"
            }
            
            return jsonify({
                "success": True,
                "document_id": doc_id,
                "message": "File processed successfully"
            })
        else:
            return jsonify({"error": "Failed to process file"}), 500
            
    except Exception as e:
        logger.error(f"Upload error: {e}")
        return jsonify({"error": "Upload failed"}), 500

@app.route('/upload-drive', methods=['POST'])
def upload_from_drive():
    """Handle Google Drive upload"""
    try:
        data = request.json
        drive_url = data.get('drive_url')
        
        if not drive_url:
            return jsonify({"error": "No Drive URL provided"}), 400
        
        # Extract folder ID from URL
        folder_id = extract_folder_id(drive_url)
        if not folder_id:
            return jsonify({"error": "Invalid Google Drive URL"}), 400
        
        # Process Google Drive files
        processed_content = process_drive_folder(folder_id)
        
        if processed_content:
            doc_id = str(uuid.uuid4())
            document_store[doc_id] = {
                "content": processed_content,
                "source": "google_drive",
                "folder_id": folder_id,
                "upload_time": datetime.now().isoformat(),
                "type": "drive_folder"
            }
            
            return jsonify({
                "success": True,
                "document_id": doc_id,
                "message": "Google Drive files processed successfully"
            })
        else:
            return jsonify({"error": "Failed to process Google Drive files"}), 500
            
    except Exception as e:
        logger.error(f"Drive upload error: {e}")
        return jsonify({"error": "Google Drive upload failed"}), 500
@app.route('/chat', methods=['POST'])
def chat():
    """Handle chat queries for multiple documents"""
    try:
        data = request.json
        query = data.get('query', '').strip()
        document_ids = data.get('document_ids', [])

        if not query:
            return jsonify({"error": "No query provided"}), 400

        if not document_ids:
            return jsonify({"error": "No document IDs provided"}), 400

        # Combine contents of selected documents
        combined_content = ""
        for doc_id in document_ids:
            if doc_id in document_store:
                combined_content += document_store[doc_id]['content'] + "\n"
            else:
                return jsonify({"error": f"Document ID {doc_id} not found"}), 404

        # Analyze with Gemini
        response = gemini_analyzer.analyze_resume(combined_content, query)

        return jsonify({
            "success": True,
            "response": response,
            "query": query,
            "timestamp": datetime.now().isoformat()
        })

    except Exception as e:
        logger.error(f"Chat error: {e}")
        return jsonify({"error": "Analysis failed"}), 500


def format_human_readable(summary: dict) -> str:
    try:
        lines = []

        lines.append(f"**Name:** {summary.get('Name', 'N/A')}")
        lines.append(f"**Email:** {summary.get('Email', 'N/A')}")
        lines.append(f"**Phone:** {summary.get('Phone', 'N/A')}")
        lines.append(f"**Years of Experience:** {summary.get('Years of Experience (estimate)', 'N/A')}")
        lines.append("\n### Key Skills:")
        for skill in summary.get("Key Skills", []):
            lines.append(f"- {skill}")
        lines.append("\n### Education:")
        for edu in summary.get("Education", []):
            lines.append(f"- {edu.get('Degree', 'N/A')} at {edu.get('Institution', 'N/A')} ({edu.get('Years', 'N/A')}) - GPA: {edu.get('GPA', 'N/A')}")
        lines.append(f"\n**Current/Latest Job Title:** {summary.get('Current/Latest Job Title', 'N/A')}")
        lines.append("\n### Companies Worked At:")
        for company in summary.get("Companies Worked At", []):
            lines.append(f"- {company}")
        lines.append("\n### Key Achievements:")
        for achievement in summary.get("Key Achievements", []):
            lines.append(f"- {achievement}")
        
        return "\n".join(lines)
    except Exception as e:
        return "Could not generate readable summary."

@app.route('/summary/<doc_id>', methods=['GET'])
def get_summary(doc_id):
    try:
        if doc_id not in document_store:
            return jsonify({"error": "Document not found"}), 404

        document = document_store[doc_id]
        content = document['content']
        
        # Extract summary
        summary = gemini_analyzer.extract_resume_summary(content)

        return jsonify({
            "success": True,
            "summary_json": summary,
            "summary_text": format_human_readable(summary),
            "document_info": {
                "type": document['type'],
                "upload_time": document['upload_time']
            }
        })
    except Exception as e:
        logger.error(f"Summary error: {e}")
        return jsonify({"error": "Failed to generate summary"}), 500

def process_single_file(file_path: str) -> str:
    """Process a single uploaded file"""
    try:
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            # Convert PDF to DOCX first
            docx_path = file_path.replace('.pdf', '.docx')
            if doc_processor.pdf_to_docx(file_path, docx_path):
                content = doc_processor.extract_text_from_docx(docx_path)
                os.remove(docx_path)  # Clean up
            else:
                # Fallback to direct text extraction
                content = doc_processor.extract_text_from_pdf(file_path)
        
        elif file_ext in ['.docx', '.doc']:
            content = doc_processor.extract_text_from_docx(file_path)
        
        else:
            logger.error(f"Unsupported file type: {file_ext}")
            return None
        
        # Clean up uploaded file
        os.remove(file_path)
        
        return content
        
    except Exception as e:
        logger.error(f"Error processing file: {e}")
        return None

def process_drive_folder(folder_id: str) -> str:
    """Process all files from Google Drive folder"""
    try:
        # Note: This is a simplified version
        # In production, you'd need proper OAuth flow
        logger.info(f"Processing Drive folder: {folder_id}")
        
        # For now, return a placeholder
        # In full implementation, you'd:
        # 1. Authenticate with Google Drive
        # 2. List all files in folder
        # 3. Download each file
        # 4. Convert PDFs to DOCX
        # 5. Combine all content
        
        return "Google Drive integration requires proper OAuth setup. Please upload files directly for now."
        
    except Exception as e:
        logger.error(f"Error processing Drive folder: {e}")
        return None

def extract_folder_id(drive_url: str) -> str:
    """Extract folder ID from Google Drive URL"""
    try:
        # Pattern for Google Drive folder URLs
        patterns = [
            r'folders/([a-zA-Z0-9-_]+)',
            r'id=([a-zA-Z0-9-_]+)',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, drive_url)
            if match:
                return match.group(1)
        
        return None
        
    except Exception as e:
        logger.error(f"Error extracting folder ID: {e}")
        return None

@app.route('/documents', methods=['GET'])
def list_documents():
    """List all uploaded documents"""
    try:
        docs = []
        for doc_id, doc_info in document_store.items():
            docs.append({
                "id": doc_id,
                "type": doc_info['type'],
                "upload_time": doc_info['upload_time'],
                "filename": doc_info.get('filename', 'N/A')
            })
        
        return jsonify({
            "success": True,
            "documents": docs,
            "count": len(docs)
        })
        
    except Exception as e:
        logger.error(f"Error listing documents: {e}")
        return jsonify({"error": "Failed to list documents"}), 500

@app.route('/delete/<doc_id>', methods=['DELETE'])
def delete_document(doc_id):
    """Delete a document from storage"""
    try:
        if doc_id in document_store:
            del document_store[doc_id]
            return jsonify({"success": True, "message": "Document deleted"})
        else:
            return jsonify({"error": "Document not found"}), 404
            
    except Exception as e:
        logger.error(f"Error deleting document: {e}")
        return jsonify({"error": "Failed to delete document"}), 500

if __name__ == '__main__':
    # Check for required environment variables
    required_vars = ['GEMINI_API_KEY']
    missing_vars = [var for var in required_vars if not os.environ.get(var)]
    
    if missing_vars:
        logger.warning(f"Missing environment variables: {missing_vars}")
        logger.warning("Please set these variables for full functionality")
    
    app.run(debug=True, host='0.0.0.0', port=5000)